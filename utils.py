import streamlit as st
import pandas as pd
import requests
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import io

def load_local_css(file_name):
    """Loads a local CSS file into the Streamlit app."""
    try:
        with open(file_name, "r") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning(f"CSS file not found: {file_name}")

def parse_csv(uploaded_file):
    """Parses an uploaded CSV file and returns a pandas DataFrame."""
    try:
        return pd.read_csv(uploaded_file)
    except Exception as e:
        st.error(f"Error parsing CSV file: {e}")
        return pd.DataFrame()

def call_ai(prompt):
    """Calls the Gemini AI with a prompt and returns the response."""
    try:
        # The application will read the Gemini API key from your secrets.toml file
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"An error occurred with the Gemini API: {e}")
        return f"Error: {e}"

def generate_docx(title, content):
    """Generates a .docx file from markdown-like text content."""
    doc = Document()
    doc.add_heading(title, level=1)
    
    for line in content.split('\n'):
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('* '): doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
        elif line.strip(): doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()